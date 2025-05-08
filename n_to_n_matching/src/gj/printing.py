#!/usr/bin/env python

# Copyright 2025 Kinu Garage
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import datetime
import docx
from docx.shared import Pt
import logging
import os
from typing import List

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle

from gj.requirements import DateRequirement
from n_to_n_matching.workdate_player import WorkDate
from n_to_n_matching.gj_rsc_matching import GjVolunteerMatching
from n_to_n_matching.util import Util as NtonUtil
from n_to_n_matching.person_player import PersonPlayer

class GjDocx():  
    def __init__(self, output_path=os.getcwd(), logger_obj: logging.Logger=None):
        self.output_path = output_path

        if not logger_obj:
            logger_obj = self._logger = NtonUtil.get_logger(__name__)
        self._logger = logger_obj

    def _narrow_table_row_spacing(self, table, space_before=Pt(0), space_after=Pt(0)):
        """
        Narrows the spacing before and after paragraphs within all cells of a table.

        Args:
            table: The docx.table.Table object.
            space_before: The spacing before each paragraph in points (docx.shared.Pt).
            space_after: The spacing after each paragraph in points (docx.shared.Pt).
        @ref: This method was suggested by Google Search
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = space_before
                    paragraph_format.space_after = space_after

    def _print_table_row(self, table_obj, row_id, person_objs: List[PersonPlayer], date: WorkDate, role: str= "委員", increment_at_end=True):
        """
        @param increment_at_end: If False, do not increment at the very end of the loop.
        """
        for person in person_objs:
            row_obj = table_obj.rows[row_id].cells
            row_obj[0].text = date
            row_obj[2].text = role
            row_obj[3].text = str(person.grade_class)
            row_obj[4].text = str(person)
            row_obj[5].text = person.phone_num
            row_id += 1
            self._logger.debug(f"grade-class: {person.grade_class}")
        if not increment_at_end:
            row_id -= 1
        return row_id
        
    def print_distributable(
            self,
            solution: GjVolunteerMatching,
            requirements: DateRequirement,
            heading1: str,
            paragraph: str='This is a sample paragraph.',
            timestamp: str="",
            path_input_file=""):

        # TODO This might need to be flexible
        TABLE_TOP_ROW = ["日付", "順", "担当", "学級", "生徒氏名", "電話番号"]

        # `_table_header_length` is 1 because the header row is the one contains`TABLE_TOP_ROW` and no other rows.
        # TODO This might have to be flexible in the future.
        _table_header_length = 1

        document = docx.Document()
        document.add_heading(heading1, level=1)
        document.add_paragraph(paragraph)

        # Total num of assignees over the all dates.
        _all_dates = solution.dates_lgtm  # `solution.dates_failed` should not be included, right?
        _total_num_assignees_alldates = 0
        for date in _all_dates:
            _day_total_assignees_num = date.total_assignees_num()
            _total_num_assignees_alldates += _day_total_assignees_num
            self._logger.debug(f"Adding {_day_total_assignees_num} assignees from {date=}")
        self._logger.info(f"{_total_num_assignees_alldates=}")

        table = document.add_table(
            rows=(_total_num_assignees_alldates + _table_header_length),
            cols=len(TABLE_TOP_ROW))

        # Adjusting column width. 
        table.autofit = True
        # TODO These are too adhoc, needs to be configurable.
        table.columns[0].width = docx.shared.Cm(3)
        table.columns[1].width = docx.shared.Cm(1)
        table.columns[4].width = docx.shared.Cm(5)
        table.columns[5].width = docx.shared.Cm(5)

        # Title row in the table
        hdr_cells = table.rows[0].cells
        for index, col_title in enumerate(TABLE_TOP_ROW, 0):
            # TODO This logic assumes the order of the elements in `TABLE_TOP_ROW` is 
            hdr_cells[index].text = col_title

        _row_id = _table_header_length  # Rows to print persons start after the heaer row.
        for date, date_detail in solution.items():
            leader_p_objs: List[PersonPlayer] = date_detail[WorkDate.ATTR_LIST_ASSIGNED_LEADER]
            committee_p_objs: List[PersonPlayer] = date_detail[WorkDate.ATTR_LIST_ASSIGNED_COMMITTEE]
            general_p_objs: List[PersonPlayer] = date_detail[WorkDate.ATTR_LIST_ASSIGNED_GENERAL]
            _beginning_row_id = _row_id
            _row_id = self._print_table_row(table, _row_id, leader_p_objs, date=date, role="リーダー")
            self._logger.debug(f"{date=}. AFTER leader: {_row_id=}")

            _row_id = self._print_table_row(table, _row_id, committee_p_objs, date=date)
            self._logger.debug(f"{date=}. AFTER commitee: {_row_id=}")

            _row_id = self._print_table_row(table, _row_id, general_p_objs, date=date, role="保護者", increment_at_end=False)
            self._logger.info(f"{date=}. AFTER general: {_row_id=}")

            # Merging ID column for the date.
            _ending_row_id = _row_id
            self._logger.info(f"{_beginning_row_id=}. {_ending_row_id=}")

            # TODO Line below is an attempt to merge cells that contain the same `date` value.
            # For some reasons this results in undesired look of the table so needs more work.
            #table.cell(_beginning_row_id, 0).merge(table.cell(_ending_row_id, 0))

            _row_id += 1

        self._narrow_table_row_spacing(table, space_before=Pt(0), space_after=Pt(0))

        # Pring the input file name
        if path_input_file:
            document.add_paragraph("使用マスタファイル名：" + os.path.basename(path_input_file))

        _sys_timestamp = datetime.datetime.strftime(datetime.datetime.now(), '%Y-%m-%d-%H-%M-%S')
        if not timestamp:
            timestamp = _sys_timestamp
        document.add_paragraph(timestamp + " 更新")

        document.add_page_break()
        document.save(os.path.join(self.output_path, f"{_sys_timestamp}_GJLS_{requirements.type_duty.name}.docx"))
